# Standard library imports
import copy
import getpass
import logging
import os
import re
import tomllib
from pathlib import Path
from typing import List, Optional, Set, Tuple

# Related third-party imports
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph
from docx2pdf import convert


# Every key contains the row no, column no and (later) text to be added to the report table cell
original_data = {
    "id": [1, 0],
    "project":[1, 1],
    "category":[1, 2],
    "view-status":[1, 3],
    "date-submitted":[1, 4],
    "last-modified":[1, 5],
    "reporter":[3, 1],
    "assigned-to":[3, 3],
    "priority":[4, 1],
    "severity":[4, 3],
    "reproducibility":[4, 5],
    "status":[5, 1],
    "resolution":[5, 3],
    "summary":[7, 1],
    "description":[8, 1],
    "steps-to-reproduce":[9, 1],
    "custom-field":{
        0: [11, 1],
        1: [12, 1],
        2: [13, 1],
        3: [14, 1],
        4: [15, 1],
        5: [16, 1],
        6: [17, 1],
        7: [18, 1],
        8: [19, 1],
        9: [20, 1],
        10: [21, 1],
        11: [22, 1],
    },
}


# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()
ENV = os.getenv("APP_ENV", "production")

# Constants
BASE_URL = os.getenv("BASE_URL") or logger.error("BASE_URL not set in .env file") or exit(1)
USERNAME_URL = os.getenv("USERNAME_URL") or logger.error("USERNAME_URL not set in .env file") or exit(1)
PASSWORD_URL = os.getenv("PASSWORD_URL") or logger.error("PASSWORD_URL not set in .env file") or exit(1)

if ENV == "debug":
    APP_USERNAME = os.getenv("APP_USERNAME") or logger.error("APP_USERNAME not set in .env file") or exit(1)
    APP_PASSWORD = os.getenv("APP_PASSWORD") or logger.error("APP_PASSWORD not set in .env file") or exit(1)

REPORT_DIR = Path("reports")
ATTACHEMENT_DIR = Path("attachements")
ISSUE_FILE = Path("issue_list.toml")


class MantisScraper:
    """A class to scrape issues and download files from a MantisBT instance."""
    
    def __init__(self, base_url: str, username_url: str, password_url: str, username: str, password: str):
        self.base_url = base_url.rstrip('/')
        self.username_url = username_url
        self.password_url = password_url
        self.username = username
        self.password = password
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.session.close()
        logger.info("Session closed.\n")
    
    def add_hyperlink(self, paragraph: Paragraph, url: str, text: str):
        """Add a hyperlink to a paragraph in a Word document.

        Args:
            paragraph (Paragraph): The paragraph to add the hyperlink to (e.g., in a table cell or document body).
            url (str): The URL for the hyperlink.
            text (str): The display text for the hyperlink.
        """
        # Add a relationship for the hyperlink
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        # Create the w:hyperlink tag
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        # Create a run element for the hyperlink text
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Add hyperlink style (blue, underlined)
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)

        # Add the text to the run
        t = OxmlElement('w:t')
        t.text = text
        run.append(rPr)
        run.append(t)
        hyperlink.append(run)

        # Add the hyperlink to the paragraph
        paragraph._element.append(hyperlink)

    def get_issue_list(self, file_path: Path) -> List[str]:
        """Read issue numbers from a file."""
        try:
            if not file_path.is_file():
                raise FileNotFoundError(f"Issue list file {file_path} does not exist.")
            with file_path.open("rb") as f:
                data = tomllib.load(f)
                issues = data["active_issues"]
            logger.info(f"Loaded {len(issues)} issues from {file_path}\n")
            return issues
        except Exception as e:
            logger.error(f"Failed to read issue list: {e}\n")
            raise

    def login(self) -> None:
        """Log in to the MantisBT instance."""
        try:
            # Step 1: Submit username
            username_payload = {"return": "index.php", "username": self.username}
            response = self.session.post(self.username_url, data=username_payload, timeout=10)
            response.raise_for_status()
            if f"Enter password for '{self.username}'" not in response.text:
                raise ValueError("Username submission failed.")

            # Step 2: Submit password
            password_payload = {"return": "login.php", "username": self.username, "password": self.password}
            response = self.session.post(self.password_url, data=password_payload, timeout=10)
            response.raise_for_status()
            if "Assigned to Me (Unresolved)" not in response.text:
                raise ValueError("Login unsuccessful.")
            logger.info("Login successful.\n")
        except requests.RequestException as e:
            logger.error(f"Login failed: {e}\n")
            raise

    def access_issue_page(self, issue: str) -> Optional[requests.Response]:
        """Access a protected issue page."""
        try:
            issue_url = f"{self.base_url}/view.php?id={issue}"
            response = self.session.get(issue_url, timeout=10)
            response.raise_for_status()
            if "View Issue Details" not in response.text:
                logger.error(f"Failed to access issue page {issue}")
                return None
            print("")
            logger.info(f"Successfully accessed issue page {issue}")
            return response, issue_url
        except requests.RequestException as e:
            logger.error(f"Error accessing issue {issue}: {e}\n")
            return None

    def scrape_page(self, response: requests.Response, issue: str) -> Tuple[BeautifulSoup, str, Path]:
        """Scrape issue page content and save as HTML."""
        try:
            soup = BeautifulSoup(response.text, "html.parser")
            issue_no = soup.title.text.split(':')[0]
            
            _category_no = "".join(soup.find("td", class_="bug-category").text.split())
            category_no = re.sub(r'[<>:"/\\|?*]', '_', f"{_category_no}_({issue_no})")

            REPORT_DIR.mkdir(parents=True, exist_ok=True)

            report_path = REPORT_DIR / category_no
            report_path.mkdir(parents=True, exist_ok=True)

            # Save HTML report
            with (report_path / f"{issue_no}_report.html").open("w", encoding="utf-8") as f:
                f.write(soup.prettify())
            logger.info(f"Saved HTML report for issue {issue_no}")
            return soup, report_path
        except Exception as e:
            logger.error(f"Error scraping issue {issue}: {e}")
            raise
    
    def get_report_data(self, issue_data:dict, soup: BeautifulSoup) -> None:
        
        for key in issue_data.keys():
            if key == "custom-field":
                custom_tags = soup.find_all(lambda tag: tag.name == 'td' and tag.get('class') == ['bug-custom-field'])
                for index, tag in enumerate(custom_tags):
                    tag_text = tag.text.strip()
                    issue_data[key][index].append(tag_text)
            else:
                tag = soup.find_all(lambda tag: tag.has_attr("class") and tag["class"] == [f"bug-{key}"])[-1]
                tag_text = tag.text.strip()
                issue_data[key].append(tag_text)

    def populate_report(self, issue_data:dict, report_path: Path, issue_url: str):
        document = Document('hrc_report_template.docx')

        style = document.styles['Normal']
        font = style.font
        font.name = "Aptos (Body)"
        font.size = Pt(8)

        _table = (document.tables)[0]

        for key in issue_data.keys():
            if key == "id":
                # Add id cell data with hyperlink
                id_cell = _table.cell(issue_data["id"][0], issue_data["id"][1])
                id_cell_text = issue_data["id"][2]
                id_cell_paragraph = id_cell.paragraphs[0]
                self.add_hyperlink(id_cell_paragraph, issue_url, id_cell_text)
            elif key == "custom-field":
                for index, value in enumerate(key):
                   _table.cell(issue_data[key][index][0], issue_data[key][index][1]).text = issue_data[key][index][2]
            else:
                _table.cell(issue_data[key][0], issue_data[key][1]).text = issue_data[key][2]
        
        # Style table row height
        for index, row in enumerate(_table.rows):
            if index <= 7 or index >= 10:
                row.height = Cm(0.5)
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        
        for root, dirs, files in os.walk(report_path / ATTACHEMENT_DIR):
            for filename in files:
                _par = document.add_paragraph()
                self.add_hyperlink(_par, f"{ATTACHEMENT_DIR}/{filename}", filename)
        
        document_file_name = f"{issue_data["category"][2]}-{issue_data["custom-field"][11][2]}.docx"
        document.save(os.path.join(report_path, document_file_name))

        # Save to PDF
        convert(os.path.join(report_path, document_file_name))
    
    def get_unique_links(self, file_links: List[BeautifulSoup]) -> Set[BeautifulSoup]:
        """Ensure downloadable file links are unique with no duplicates

        Args:
            file_links (List[BeautifulSoup]): A list of BeautifulSoup 4 element tags

        Returns:
            Set[BeautifulSoup]: A unique set of BeautifulSoup4 element tags
        """
        try:
            links = [link for link in file_links if len(link.attrs) == 1 and link.find() is None and link.get_text().strip()]
            unique_links = set(links)
            return unique_links
        except Exception as e:
            logger.error(e)
            raise

    def download_multiple_type_files(self, report_path: Path, file_links: Set[BeautifulSoup]) -> None:
        """Download files from provided links."""
        if not file_links:
            logger.warning("No file links found.")
            return

        logger.info(f"Found {len(file_links)} file links.")

        for index, link in enumerate(file_links, 1):
            try:
                relative_url = link["href"]
                file_url = relative_url if relative_url.startswith('http') else f"{self.base_url}/{relative_url.lstrip('/')}"

                file_basename, _file_extension = os.path.splitext(link.text)
                file_extension = _file_extension.lstrip('.')
                
                # This regex sanitizes the file name text:
                    # input:    My new:file/name.pdf
                    # output:   My_new_file_name.pdf
                filename = re.sub(r'[^\w\.-]', '_', f"{index}_{file_basename.strip()}.{file_extension}" or f"{index}_file.{file_extension}")
                (report_path / ATTACHEMENT_DIR).mkdir(parents=True, exist_ok=True)
                file_path = report_path / ATTACHEMENT_DIR / filename

                logger.info(f"Downloading {filename} from {file_url}")
                response = self.session.get(file_url, timeout=10, allow_redirects=True)
                response.raise_for_status()

                with file_path.open("wb") as f:
                    f.write(response.content)
                logger.info(f"Saved {filename}")
            except requests.RequestException as e:
                logger.error(f"Failed to download {file_url}: {e}")

def main():
    """Main function to orchestrate the scraping process."""
    try:
        if ENV == "debug":
            # === Use in dev enironment only ===
            username = APP_USERNAME
            password = APP_PASSWORD
            # === Use in dev enironment only ===

        elif ENV == "production":
            # Prompt for credentials
            username = input("Enter your Mantis username: ")
            password = getpass.getpass(prompt=f"Enter password for {username}: ")

        # Ensure downloads directory exists
        REPORT_DIR.mkdir(parents=True, exist_ok=True)

        with MantisScraper(BASE_URL, USERNAME_URL, PASSWORD_URL, username, password) as scraper:
            scraper.login()
            issues = scraper.get_issue_list(ISSUE_FILE)
            for issue in issues:
                response, issue_url = scraper.access_issue_page(issue)
                if response:
                    # Make a fresh copy of the original dict for this iteration
                    issue_data = copy.deepcopy(original_data)

                    soup, report_path = scraper.scrape_page(response, issue)
                    file_links = soup.find_all("a", href=lambda x: x and "file_download.php" in x)
                    
                    unique_links = scraper.get_unique_links(file_links)
                    scraper.download_multiple_type_files(report_path, unique_links)

                    scraper.get_report_data(issue_data, soup)
                    scraper.populate_report(issue_data, report_path, issue_url)
    except Exception as e:
        logger.error(f"Script failed: {e}")
        raise

if __name__ == "__main__":
    main()